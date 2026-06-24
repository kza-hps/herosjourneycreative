#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Fix the 6 glossary entries that didn't get pronunciations due to special characters."""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

from docx import Document

GLOSSARY_PATH = 'public/journal/Ho and the Baby Eater - Glossary Publication.docx'

# Exact term bytes → pronunciation
# Determined by inspecting cell bytes:
#   ʻ  = U+02BB  (modifier letter turned comma, Hawaiian okina)
#   '  = U+2018  (left single quotation mark)
#   '  = U+0027  (straight apostrophe)

FIXES = {
    "Ho's atoll":                 "HOH's ah-TOL",       # U+0027 straight apostrophe
    "Na-Mala-o-Kala‘i":           "nah-MAH-lah-oh-kah-LAH-ee",  # U+2018
    "Hōkūleʻa":         "HOH-koo-LEH-ah",     # Hōkūleʻa
    "Hōkūpaʻa":         "HOH-koo-PAH-ah",     # Hōkūpaʻa
    "Paiʻea":                     "pai-EH-ah",           # Paiʻea
    "Tuʻunaga":                   "too-oo-NAH-gah",      # Tuʻunaga
}

def main():
    doc = Document(GLOSSARY_PATH)
    fixed = 0
    for table in doc.tables:
        for row in table.rows[1:]:
            term = row.cells[0].text.strip()
            if term in FIXES:
                cell = row.cells[0]
                p = cell.paragraphs[0]
                p.add_run(f' ({FIXES[term]})')
                print(f'  Fixed: {repr(term)}')
                fixed += 1

    print(f'\n{fixed} of {len(FIXES)} terms fixed.')
    doc.save(GLOSSARY_PATH)

main()
