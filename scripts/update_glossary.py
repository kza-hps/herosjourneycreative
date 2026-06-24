#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Update the Ho & the Baby Eater glossary with pronunciation guides and new entries."""

import sys, io, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GLOSSARY_PATH = 'public/journal/Ho and the Baby Eater - Glossary Publication.docx'

# ============================================================
# PRONUNCIATION DATA  (keyed by exact cell text)
# ============================================================

TABLE_PRON = [
    # Table 0: Core Kupu
    {
        'āe': 'ah-EH',
        'ariki': 'ah-REE-kee',
        'atua': 'ah-TOO-ah',
        'auē': 'ow-EH',
        'hā': 'HAH',
        'hapū': 'hah-POO',
        'harakeke': 'hah-rah-KEH-keh',
        'hoe': 'HOH-eh',
        'kaihautū': 'kai-how-TOO',
        'karaka': 'kah-RAH-kah',
        'kauri': 'KOW-ree',
        'kava': 'KAH-vah',
        'kete': 'KEH-teh',
        'kōhū': 'KOH-hoo',
        'kūmara': 'KOO-mah-rah',
        'lavālava': 'lah-VAH-lah-vah',
        'mākutu': 'MAH-koo-too',
        'mana': 'MAH-nah',
        'maro': 'MAH-roh',
        'moa': 'MOH-ah',
        'motu': 'MOH-too',
        'muka': 'MOO-kah',
        'namu': 'NAH-moo',
        'ngeri': 'NGEH-ree',
        'pā': 'PAH',
        'pākē': 'PAH-keh',
        'paoa': 'pah-OH-ah',
        'patupaiarehe': 'pah-too-pai-ah-REH-heh',
        'pāua': 'PAH-oo-ah',
        'poerava': 'poh-EH-rah-vah',
        'pōhutukawa': 'POH-hoo-too-KAH-wah',
        'Pouākai': 'poh-AH-kai',
        'pounamu': 'poh-NAH-moo',
        'puaka': 'poo-AH-kah',
        'puga': 'POO-gah',
        'rāhui / rahui': 'RAH-hwee',
        'rauhā-tangi': 'row-HAH-tahng-ee',
        'rohe': 'ROH-heh',
        'tāhei': 'TAH-heh-ee',
        'tāmure': 'TAH-moo-reh',
        'tāniko': 'TAH-nee-koh',
        'taonga': 'tah-ONG-ah',
        'tapa': 'TAH-pah',
        'tapu': 'TAH-poo',
        'taro': 'TAH-roh',
        'tī': 'TEE',
        'tika': 'TEE-kah',
        'toetoe': 'TOH-eh-TOH-eh',
        'tohunga': 'toh-HUNG-ah',
        'tōtara': 'TOH-tah-rah',
        'tūrama': 'TOO-rah-mah',
        'utu': 'OO-too',
        'wairua': 'wai-ROO-ah',
        'waka': 'WAH-kah',
        'whakataukī': 'fah-kah-TOW-kee',
        'whāriki': 'FAH-ree-kee',
    },
    # Table 1: Atua
    {
        'Ārohirohi': 'AH-roh-hee-roh-hee',
        'Feke': 'FEH-keh',
        'Feke and Rapa-Iti': 'FEH-keh and RAH-pah-EE-tee',
        'Io': 'EE-oh',
        'Mailagu': 'mai-LAH-goo',
        'Mārama': 'MAH-rah-mah',
        'Māui': 'MAH-oo-ee',
        'Papatūānuku': 'pah-pah-TOO-AH-noo-koo',
        'Rā': 'RAH',
        'Ranginui': 'RANG-ee-noo-ee',
        'Rongo': 'RONG-oh',
        'Takali Foto': 'tah-KAH-lee FOH-toh',
        'Takaroa': 'tah-kah-ROH-ah',
        'Tāne / Tāne-Mahuta': 'TAH-neh / TAH-neh-mah-HOO-tah',
        'Tāwhiri / Tāwhirimātea': 'TAH-fee-ree / TAH-fee-ree-MAH-teh-ah',
        'Tokoroa': 'toh-koh-ROH-ah',
        'Tū': 'TOO',
        'Wātea': 'WAH-teh-ah',
    },
    # Table 2: Peoples / Places
    {
        'Ahukai': 'ah-HOO-kai',
        'Autara': 'ow-TAH-rah',
        'Feke': 'FEH-keh',
        "Ho’s atoll": "HOH's ah-TOL",
        'Kafiki': 'kah-FEE-kee',
        'Kafiki Motu': 'kah-FEE-kee MOH-too',
        'Lidopo': 'lee-DOH-poh',
        'Mahana': 'mah-HAH-nah',
        'Makatemā': 'mah-kah-TEH-ah',
        'Makatemān': 'mah-kah-TEH-an',
        'Makateā': 'mah-kah-TEH-ah',
        'Makateān': 'mah-kah-TEH-an',
        'Matalagi': 'mah-tah-LAH-gee',
        'Matavai': 'mah-tah-VAI',
        "Na-Mala-o-Kalaʼi": 'nah-MAH-lah-oh-kah-LAH-ee',
        'Nasara': 'nah-SAH-rah',
        'Rapa-Iti': 'RAH-pah-EE-tee',
        'Takaroans': 'tah-kah-ROH-anz',
        'Tuakau': 'too-AH-kow',
        'Tumutumu': 'too-moo-TOO-moo',
        'Ulu Waimate': 'OO-loo wai-MAH-teh',
        'Unusi': 'oo-NOO-see',
        'Vanua': 'VAH-noo-ah',
        'Waimate River': 'wai-MAH-teh',
    },
    # Table 3: Characters
    {
        'Ailani': 'ai-LAH-nee',
        'Arahuta': 'ah-rah-HOO-tah',
        'Baby Eater': 'BAY-bee EE-ter',
        'Chief Kuanua': 'koo-AH-noo-ah',
        'Dark Flame': '(English)',
        'Faturaki': 'fah-too-RAH-kee',
        'Galaiga': 'gah-LAI-gah',
        'Gundaidhar': 'gun-DYE-dar',
        'Hina': 'HEE-nah',
        'Ho': 'HOH',
        'Hōkūleʼa': 'HOH-koo-LEH-ah',
        'Hōkūpaʼa': 'HOH-koo-PAH-ah',
        'Iwa': 'EE-wah',
        'Kalapa': 'kah-LAH-pah',
        'Kanopolu': 'kah-noh-POH-loo',
        'Kura': 'KOO-rah',
        'Loha': 'LOH-hah',
        'Makosai': 'mah-KOH-sai',
        'Manōkalanipō': 'mah-NOH-kah-lah-nee-POH',
        'Paiʼea': 'pai-EH-ah',
        'Pakuu': 'pah-KOO-oo',
        'Pataru': 'pah-TAH-roo',
        'Pounamu': 'poh-NAH-moo',
        'Selai': 'SEH-lai',
        'Sere': 'SEH-reh',
        'Sukey': 'SOO-kee',
        'Teā': 'TEH-ah',
        'Totokona': 'toh-toh-KOH-nah',
        'Tufukia': 'too-foo-KEE-ah',
        'Turuturu': 'too-roo-TOO-roo',
        'Tuʼunaga': 'too-oo-NAH-gah',
        'Waru': 'WAH-roo',
    },
]

# ============================================================
# NEW ENTRIES
# (table_idx, insert_after_term, term, pronunciation, meaning)
# ============================================================

NEW_ENTRIES = [
    (0, 'harakeke', 'hika', 'HEE-kah',
     'The act of fire-making by friction; rubbing a hardened pointed stick against a grooved board to generate an ember. Used in the manuscript in the phrase “rub of hika to mahoe.”'),
    (0, 'hoe', 'hua', 'HOO-ah',
     'Fruit, produce, result; in wider usage, the yield of any effort. In the manuscript used colloquially for brains or mind: “this cursed island has scooped out my hua—now I’m only the shell.”'),
    (0, 'lavālava', 'mahoe', 'mah-HOH-eh',
     'A fast-growing native tree whose soft wood traditionally serves as the base board in hika fire-making.'),
    (0, 'maro', 'miro', 'MEE-roh',
     'A native forest tree bearing bright red berries; appears in jungle and forest-path descriptions in the manuscript.'),
    (0, 'moa', 'moai', 'MOH-ai',
     'The monumental stone statues of Rapa Nui (Easter Island); used in the manuscript as a simile for crushing, immovable weight.'),
    (0, 'moai', 'moi', 'MOH-ee',
     'Threadfin fish; a valued Pacific food fish. The first named creature in the novel, dying on the end of Ho’s spear in the opening scene.'),
    (0, 'ngeri', 'ōtaha', 'OH-tah-hah',
     'Frigate bird; a large seabird that soars effortlessly on thermals. In the manuscript one is struck from the air by a falling stone egg during the sky-tear event.'),
    (0, 'paoa', 'paru', 'PAH-roo',
     'Mud, mire; ruined or not working. Used in the manuscript as colloquial slang for something broken or out of use: “Patu gone paru.”'),
    (0, 'paru', 'patu', 'PAH-too',
     'Short-handled striking weapon; a Māori hand-club used in close combat. Also used in colloquial and literary contexts as a euphemism for the penis.'),
    (0, 'rāhui / rahui', 'rangatahi', 'rang-ah-TAH-hee',
     'Young people; youth. Used as a descriptor in the manuscript: atua rangatahi, young divine offspring of mixed descent.'),
    (0, 'tāniko', 'taniwha', 'tah-NEE-fah',
     'Supernatural creature associated with deep water; powerful and often dangerous beings that inhabit rivers, lakes, and the sea.'),
    (0, 'utu', 'vokai', 'VOH-kai',
     'The Fijian Vokai iguana; a critically endangered tree-dwelling lizard native to Fiji, known for vibrant emerald green colouring with white and black banding. Can darken its skin to near-black when stressed or threatened.'),
    (0, 'waka', 'whakapapa', 'fah-kah-PAH-pah',
     'Genealogy; the layering of one generation upon another. Recitation of ancestral lineage establishes identity, standing, and spiritual connection. Ho’s uncertainty about his whakapapa is central to his crisis of belonging.'),
    # Atua
    (1, 'Feke and Rapa-Iti', 'Haronga', 'hah-RONG-ah',
     'Dawn god; atua of warmth and the first light of day. In the manuscript, Ho wonders whether Haronga has opened his eyes to witness the tears in the sky—unsettled to be watched by a god he had barely considered.'),
    # Peoples / Places
    (2, 'Kafiki Motu', 'Lapita', 'lah-PEE-tah',
     'The Lapita people; ancient seafarers whose distinctive pottery tradition marks the earliest known settlement of western Polynesia, c. 1600–500 BCE. Referenced in the manuscript as a historical antecedent to the world of Kafiki.'),
]

# ============================================================
# PRONUNCIATION GUIDE CONTENT
# ============================================================

PRON_GUIDE_HEADING = 'Pronunciation Guide'

PRON_GUIDE_PARAS = [
    ('normal',
     'Māori and most Polynesian languages share a consistent vowel system. Once the sounds below are learned, any word in this glossary can be sounded out reliably. The pronunciations shown in brackets after each entry use a plain-English respelling: syllables are separated by hyphens and the stressed syllable appears in CAPS.'),
    ('heading2', 'Vowels'),
    ('normal',
     'Five pure sounds, short unless marked with a macron:'),
    ('normal',
     '  a — as in father        e — as in bed        i — as in feet (short)'),
    ('normal',
     '  o — as in core         u — as in moon (short)'),
    ('heading2', 'Macrons (ā ē ī ō ū)'),
    ('normal',
     'The same five vowels held longer. A macron roughly doubles the duration of the vowel without changing its quality: ā is a held “ah,” ō is a held “oh.”'),
    ('heading2', 'Diphthongs'),
    ('normal',
     'Two vowels blending smoothly together without a break:'),
    ('normal',
     '  ae → ah-eh      ai → ah-ee      ao → ah-oh      au → ah-oo'),
    ('normal',
     '  ei → eh-ee      oe → oh-eh      oi → oh-ee      ou → oh-oo'),
    ('heading2', 'Consonants'),
    ('normal',
     'Mostly close to English, with three key differences:'),
    ('normal',
     '  wh — pronounced as f in most Māori dialects (e.g., whakapapa → fah-kah-PAH-pah; whāriki → FAH-ree-kee)'),
    ('normal',
     '  ng — always the sound in singing, never as in finger; it can begin a syllable (e.g., ngeri begins with this nasal)'),
    ('normal',
     '  r — a light flap, closer to the r in Spanish pero than the English r'),
    ('heading2', 'Stress'),
    ('normal',
     'Generally falls on the first syllable. Macronned vowels draw additional weight wherever they appear.'),
]

# ============================================================
# HELPERS
# ============================================================

def make_t(text):
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return t

def make_run(text):
    r = OxmlElement('w:r')
    r.append(make_t(text))
    return r

def make_para_xml(text, style_id=None):
    p = OxmlElement('w:p')
    if style_id:
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style_id)
        pPr.append(pStyle)
        p.append(pPr)
    p.append(make_run(text))
    return p

def get_heading_style_id(doc, level):
    name = f'Heading {level}'
    for style in doc.styles:
        if style.name == name:
            return style.style_id
    return f'Heading{level}'

def get_base_term(text):
    """Strip pronunciation from term cell text for matching."""
    text = text.strip()
    idx = text.rfind(' (')
    if idx >= 0:
        return text[:idx]
    return text

def add_pron_to_row(row, pron):
    """Append ' (pron)' to first cell's first paragraph."""
    cell = row.cells[0]
    p = cell.paragraphs[0]
    p.add_run(f' ({pron})')

def clone_row(template_row, term, pron, meaning):
    new_tr = copy.deepcopy(template_row._tr)
    tcs = new_tr.findall('.//' + qn('w:tc'))

    for tc, content in zip(tcs[:2], [f'{term} ({pron})', meaning]):
        for p in tc.findall('.//' + qn('w:p')):
            # Remove all runs and hyperlinks, keep pPr
            for child in list(p):
                if child.tag not in [qn('w:pPr')]:
                    p.remove(child)
            p.append(make_run(content))

    return new_tr

def insert_after(table, after_term, term, pron, meaning):
    tmpl = table.rows[1]
    new_tr = clone_row(tmpl, term, pron, meaning)
    for row in table.rows:
        if get_base_term(row.cells[0].text) == after_term:
            row._tr.addnext(new_tr)
            return True
    print(f'  WARNING: after_term {after_term!r} not found in table', file=sys.stderr)
    return False

# ============================================================
# MAIN
# ============================================================

def main():
    doc = Document(GLOSSARY_PATH)

    # ── 1. Update document title ──────────────────────────────
    title_para = doc.paragraphs[0]
    for run in title_para.runs:
        if 'Glossary' in run.text and 'Pronunciation' not in run.text:
            run.text = run.text.replace('Glossary', 'Glossary and Pronunciation Guide')
            print(f'  Title updated: {run.text!r}')
            break

    # ── 2. Insert pronunciation guide before Core Kupu heading ─
    # paragraphs[7] should be the Core Kupu heading
    core_kupu_para = None
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1' and 'Core Kupu' in p.text:
            core_kupu_para = p
            break
    if core_kupu_para is None:
        print('ERROR: Could not find Core Kupu heading', file=sys.stderr)
        sys.exit(1)

    h1_id = get_heading_style_id(doc, 1)
    h2_id = get_heading_style_id(doc, 2)

    # Insert guide paragraphs in reverse order so they land in correct order
    for (kind, text) in reversed(PRON_GUIDE_PARAS):
        style_id = h2_id if kind == 'heading2' else None
        p_xml = make_para_xml(text, style_id=style_id)
        core_kupu_para._p.addprevious(p_xml)

    # Insert the Pronunciation Guide heading
    guide_heading_xml = make_para_xml(PRON_GUIDE_HEADING, style_id=h1_id)
    core_kupu_para._p.addprevious(guide_heading_xml)
    print('  Pronunciation guide section inserted.')

    # ── 3. Add pronunciations to existing rows ────────────────
    unmatched = []
    for ti, table in enumerate(doc.tables):
        if ti >= len(TABLE_PRON):
            break
        pron_map = TABLE_PRON[ti]
        for row in table.rows[1:]:
            term = row.cells[0].text.strip()
            if term in pron_map:
                add_pron_to_row(row, pron_map[term])
            else:
                unmatched.append((ti, term))

    if unmatched:
        print('\n  Unmatched terms (no pronunciation added):', file=sys.stderr)
        for ti, t in unmatched:
            print(f'    Table {ti}: {t!r}', file=sys.stderr)

    # ── 4. Insert new entries ─────────────────────────────────
    for (ti, after_term, term, pron, meaning) in NEW_ENTRIES:
        ok = insert_after(doc.tables[ti], after_term, term, pron, meaning)
        status = 'OK' if ok else 'FAILED'
        print(f'  [{status}] Insert {term!r} after {after_term!r} in table {ti}')

    # ── 5. Save ───────────────────────────────────────────────
    doc.save(GLOSSARY_PATH)
    print('\nDone. Glossary saved.')

main()
